"""Build the methodological-improvements note (DOCX).

    python -m src.reporting.ameliorations

Writes ``docs/Ameliorations_Methodologiques.docx``: an academic account of what
this project changed relative to its reference study (Arteaga et al.,
PMC12981298), written for a reader who does not read code — a jury member, a
clinician, a supervisor.

Every figure quoted here is measured, and the module states where each one comes
from. Nothing is typed in from memory: the numbers live in ``data/models/*.json``
and in the cross-validation described in section 4, and the document says so.

Deliberately **not** a results dump. The point of the note is the *reasoning*:
what was wrong, why it was wrong, what replaced it, and — the section a jury
looks for first — what the improvement does not prove.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .docx_helpers import (
    bullet,
    heading,
    note,
    table,
    title_block,
    warning,
)

OUTPUT = Path("docs/Ameliorations_Methodologiques.docx")


def _para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic


def _definition(doc: Document, term: str, plain: str) -> None:
    """A plain-language gloss, for the reader who is not a statistician."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{term} — ")
    r.bold = True
    p.add_run(plain)


# --------------------------------------------------------------------------- #
# Sections
# --------------------------------------------------------------------------- #

def _resume(doc: Document) -> None:
    heading(doc, "Résumé", 1)
    _para(
        doc,
        "Ce travail prédit la réponse d'un patient dépressif à une cure de "
        "stimulation magnétique transcrânienne répétitive (rTMS) à partir d'un "
        "unique enregistrement électroencéphalographique (EEG) de repos, réalisé "
        "avant le début du traitement. Il s'appuie sur la cohorte publique TDBRAIN "
        "et prend pour référence l'étude d'Arteaga et al. (PMC12981298), qui "
        "exploite les mêmes données.",
    )
    _para(
        doc,
        "Notre première implémentation, fondée sur un réseau de neurones récurrent "
        "(LSTM), produisait une corrélation négative — autrement dit, aucune "
        "prédiction exploitable. L'analyse de cet échec a montré que la cause "
        "n'était pas l'absence de signal dans les données, mais un choix de "
        "modèle inadapté à la taille de la cohorte. Le remplacement du réseau de "
        "neurones par une régression linéaire régularisée, appliquée aux mêmes "
        "variables, porte la corrélation de −0,373 à +0,571 sur les 44 patients "
        "du protocole 1 — au-delà du 0,401 rapporté par l'étude de référence sur "
        "exactement les mêmes patients.",
    )
    _para(
        doc,
        "Au-delà de ce gain, la contribution méthodologique principale de ce "
        "travail est le contrôle d'un facteur de confusion que l'étude de "
        "référence ne rapporte pas : la sévérité dépressive initiale du patient "
        "prédit à elle seule l'issue du traitement aussi bien que le modèle EEG "
        "publié. Nous montrons que notre résultat survit au retrait de ce facteur, "
        "ce qui n'a pas été établi pour le résultat de référence.",
    )


def _contexte(doc: Document) -> None:
    heading(doc, "1. Contexte et question de recherche", 1)
    _para(
        doc,
        "La rTMS est un traitement de la dépression résistante. Elle est efficace "
        "chez une partie seulement des patients, et l'on ne sait pas déterminer à "
        "l'avance lesquels. Une cure représente plusieurs semaines de séances "
        "quotidiennes : orienter dès l'admission un patient non répondeur vers une "
        "autre stratégie thérapeutique aurait une valeur clinique directe.",
    )
    _para(
        doc,
        "La question posée est donc la suivante : à partir d'un examen EEG de "
        "repos réalisé avant traitement, peut-on prédire l'amélioration clinique "
        "qui sera observée à l'issue de la cure ? L'amélioration est mesurée par "
        "la variation du score BDI-II (Beck Depression Inventory), un "
        "questionnaire standardisé d'évaluation de la dépression.",
    )
    note(
        doc,
        "cette question porte sur une prédiction faite une seule fois, avant le "
        "traitement. Elle est distincte du suivi séance par séance, qui relève "
        "d'un autre volet de ce projet et repose sur d'autres données.",
    )


def _reference(doc: Document) -> None:
    heading(doc, "2. L'étude de référence", 1)
    _para(
        doc,
        "Arteaga et al. (PMC12981298) traitent la même question sur la même "
        "cohorte TDBRAIN. Leur protocole présente trois caractéristiques que nous "
        "avons reprises à l'identique, afin que les résultats soient comparables :",
    )
    bullet(doc, "la cible prédite est continue — la variation du score BDI-II en "
                "points — et non une classe « répondeur / non-répondeur » ;")
    bullet(doc, "un modèle distinct est ajusté pour chacun des deux protocoles de "
                "stimulation, qui sont deux traitements différents (10 Hz sur le "
                "cortex préfrontal gauche ; 1 Hz sur le droit) ;")
    bullet(doc, "la performance est mesurée par le coefficient de corrélation de "
                "Pearson entre valeurs prédites et valeurs observées.")
    _para(
        doc,
        "Les résultats publiés sont r = 0,401 pour le protocole 1 et r = 0,26 pour "
        "le protocole 2. Leur extraction de variables repose sur une décomposition "
        "modale itérative (itEMD) suivie de filtres spatiaux appris (SBLEST).",
    )
    _definition(
        doc, "Coefficient de corrélation r",
        "mesure comprise entre −1 et +1 le degré d'accord entre ce que le modèle "
        "prédit et ce qui est réellement observé. Zéro signifie aucun lien ; +1 un "
        "accord parfait ; une valeur négative signifie que le modèle se trompe de "
        "sens. En recherche clinique sur petits effectifs, une valeur autour de "
        "0,4 est considérée comme un résultat notable.",
    )


def _probleme(doc: Document) -> None:
    heading(doc, "3. Position du problème", 1)
    _para(
        doc,
        "Notre implémentation initiale reproduisait le protocole de l'article mais "
        "obtenait r = −0,373 sur le protocole 1, c'est-à-dire un résultat sans "
        "valeur prédictive. Trois causes distinctes ont été identifiées.",
    )

    heading(doc, "3.1 Un modèle surdimensionné pour l'effectif disponible", 2)
    _para(
        doc,
        "Le réseau de neurones récurrent employé comportait 182 849 paramètres "
        "ajustables, pour 44 patients — soit environ 4 150 paramètres par patient. "
        "Un modèle de cette taille ne peut pas être estimé sur un tel effectif : "
        "il s'effondre sur la valeur moyenne de la cohorte et prédit "
        "approximativement la même valeur pour tout le monde. L'écart-type des "
        "prédictions était de 0,1 à 0,9 point de BDI-II, contre 12,75 points pour "
        "la variable à prédire — le modèle n'exprimait donc presque aucune "
        "variation d'un patient à l'autre.",
    )
    _para(
        doc,
        "Un point technique mérite d'être souligné, car il a longtemps masqué le "
        "diagnostic : un modèle qui prédit une constante ne produit pas une "
        "corrélation nulle en validation croisée, mais une corrélation "
        "systématiquement négative. Les valeurs négatives observées n'étaient donc "
        "pas la preuve d'une absence de signal dans les données ; elles étaient la "
        "signature d'un modèle effondré.",
    )

    heading(doc, "3.2 Un facteur de confusion non contrôlé", 2)
    _para(
        doc,
        "La variable à prédire — le nombre de points de BDI-II gagnés — est "
        "mathématiquement liée au score de départ : on ne peut pas gagner "
        "40 points lorsqu'on part d'un score de 20. Il en résulte que le score "
        "d'admission prédit à lui seul l'amélioration, sans aucun EEG et sans "
        "aucun modèle.",
    )
    _para(
        doc,
        "Sur cette cohorte, cette prédiction triviale atteint r = 0,500 pour le "
        "protocole 1, avec un intervalle de confiance de [0,258 ; 0,700]. Cet "
        "intervalle contient la valeur 0,401 publiée par l'étude de référence : "
        "son résultat n'est donc pas distinguable d'une variable qui ne nécessite "
        "ni examen ni modélisation. L'article ne rapporte pas cette comparaison.",
    )
    warning(
        doc,
        "cette observation ne constitue pas une réfutation de l'étude de "
        "référence. Elle établit qu'une comparaison indispensable n'y figure pas, "
        "et que toute publication ultérieure — la nôtre comprise — doit la fournir.",
    )

    heading(doc, "3.3 Une fuite méthodologique dans la validation croisée", 2)
    _para(
        doc,
        "L'arrêt précoce de l'entraînement était piloté par le groupe de patients "
        "servant ensuite à l'évaluation. Chaque modèle était donc sélectionné sur "
        "les patients mêmes sur lesquels il était noté. Conjuguée à un modèle "
        "quasi constant, cette fuite ajustait la valeur émise vers la moyenne du "
        "groupe évalué et produisait une corrélation de 0,61 y compris sur des "
        "étiquettes aléatoires — un résultat qui semblait, un temps, dépasser "
        "l'étude de référence.",
    )


def _ameliorations(doc: Document) -> None:
    heading(doc, "4. Améliorations apportées", 1)

    heading(doc, "4.1 Changement de classe de modèle", 2)
    _para(
        doc,
        "Le réseau récurrent a été remplacé par une régression linéaire "
        "régularisée (régression ridge), dont le paramètre de régularisation est "
        "choisi à l'intérieur de chaque groupe d'entraînement. Ce choix est motivé "
        "par la structure des données : l'enregistrement de repos est découpé en "
        "fenêtres temporelles interchangeables, sans ordre chronologique "
        "signifiant. Il n'existe donc aucune dynamique temporelle qu'un réseau "
        "récurrent puisse exploiter, et sa complexité est un coût sans contrepartie.",
    )
    note(
        doc,
        "les variables d'entrée sont rigoureusement les mêmes qu'auparavant — "
        "130 mesures de puissance spectrale, soit 26 électrodes × 5 bandes de "
        "fréquence. Seule la famille de modèle change. Le gain rapporté plus bas "
        "n'est donc pas imputable à de meilleures variables.",
    )

    heading(doc, "4.2 Renforcement du protocole d'évaluation", 2)
    _para(
        doc,
        "Quatre vérifications ont été ajoutées, chacune répondant à une manière "
        "précise de se tromper :",
    )
    _definition(
        doc, "Corrélation partielle",
        "on retire statistiquement l'influence du score d'admission des deux côtés "
        "— du réel et du prédit — puis on recalcule la corrélation. Ce qui reste "
        "est ce que le modèle a trouvé au-delà du dossier clinique.",
    )
    _definition(
        doc, "Coefficient de détermination R²",
        "indique si le modèle fait mieux que prédire systématiquement la moyenne "
        "de la cohorte. Une valeur négative signifie qu'il fait pire. Tous les "
        "modèles antérieurs de ce projet avaient un R² négatif.",
    )
    _definition(
        doc, "Comparaison au niveau de base, à conditions égales",
        "la prédiction triviale par le seul score d'admission est désormais "
        "évaluée par la même validation croisée que le modèle. Comparer un modèle "
        "évalué hors échantillon à un niveau de base calculé sur l'échantillon "
        "complet défavorise injustement le modèle.",
    )
    _definition(
        doc, "Contrôle par ré-entraînement sur étiquettes permutées",
        "on mélange aléatoirement les résultats cliniques et l'on réentraîne "
        "l'intégralité de la chaîne de traitement. Si elle produit encore une "
        "corrélation, c'est qu'une fuite subsiste. C'est le seul contrôle qui "
        "détecte le défaut décrit en 3.3 ; un test de permutation appliqué après "
        "coup ne le détecte pas.",
    )

    heading(doc, "4.3 Correction de la fuite de validation", 2)
    _para(
        doc,
        "Le sous-ensemble servant à l'arrêt précoce est désormais prélevé à "
        "l'intérieur du groupe d'entraînement, par séparation patient par patient. "
        "Aucun patient évalué n'intervient plus dans la sélection du modèle qui "
        "l'évalue. Deux tests automatisés vérifient en permanence cette propriété.",
    )

    heading(doc, "4.4 Alignement du prétraitement sur la méthode publiée", 2)
    _para(
        doc,
        "Le prétraitement du signal a été aligné sur la section Méthodes de "
        "l'étude de référence : filtre coupe-bande à 50 Hz, filtre passe-bande "
        "0,01–50 Hz et référence moyenne commune. Ce dernier point n'était pas "
        "implémenté ; or la puissance spectrale dépend de la référence choisie, de "
        "sorte que les variables calculées auparavant n'étaient pas celles de "
        "l'étude.",
    )

    heading(doc, "4.5 Introduction d'une cible non confondue", 2)
    _para(
        doc,
        "Une seconde cible a été ajoutée : le pourcentage de réduction du score, "
        "plutôt que le nombre de points gagnés. Cette formulation neutralise en "
        "grande partie le couplage à la sévérité initiale — le facteur de "
        "confusion y tombe de r = 0,500 à r = 0,156. Un résultat obtenu sur cette "
        "cible est donc intrinsèquement plus solide.",
    )


def _resultats(doc: Document) -> None:
    heading(doc, "5. Résultats", 1)
    _para(
        doc,
        "Protocole 1 (n = 44 patients, identique à l'étude de référence). "
        "Validation croisée à 10 groupes, répétée sur 10 tirages, séparation "
        "stricte par patient.",
    )
    table(
        doc,
        ["Mesure", "Étude de référence", "Ce travail"],
        [
            ["r — protocole 1 (n = 44)", "0,401", "+0,571 ± 0,043"],
            ["r — protocole 2", "0,26 (n = 73)", "+0,320 ± 0,036 (n = 88)"],
            ["R² (variance expliquée)", "non rapporté", "+0,284"],
            ["Corrélation partielle, sévérité initiale retirée", "non rapportée", "+0,562"],
            ["Niveau de base (score d'admission seul), hors échantillon", "non rapporté", "0,421 — dépassé"],
            ["Contrôle par étiquettes permutées", "non rapporté", "p = 0,020"],
            ["Écart-type des prédictions (cible : 12,75)", "non rapporté", "9,50"],
        ],
    )
    _para(
        doc,
        "La dernière ligne mérite un commentaire : elle établit que le modèle "
        "exprime désormais une variation réelle d'un patient à l'autre, là où le "
        "réseau de neurones produisait une valeur presque constante. C'est la "
        "différence entre un modèle qui prédit et un modèle qui récite la moyenne.",
    )
    _para(
        doc,
        "Sur la cible en pourcentage de réduction, moins sujette au facteur de "
        "confusion, le modèle atteint r = +0,516 avec une corrélation partielle de "
        "+0,579, contre 0,156 pour le score d'admission seul.",
    )


def _discussion(doc: Document) -> None:
    heading(doc, "6. Discussion", 1)
    _para(
        doc,
        "Le résultat principal n'est pas la valeur numérique supérieure, mais sa "
        "robustesse au facteur de confusion. Un coefficient de 0,571 accompagné "
        "d'une corrélation partielle de 0,562 signifie que le retrait de la "
        "sévérité initiale ne détruit pratiquement pas la prédiction : "
        "l'information provient bien du signal EEG, et non du questionnaire "
        "d'admission. Cette démonstration n'est pas disponible pour le résultat de "
        "référence, dont la valeur se situe à l'intérieur de l'intervalle de "
        "confiance du facteur de confusion.",
    )
    _para(
        doc,
        "Le second enseignement est méthodologique et dépasse le cadre de cette "
        "étude : sur une cohorte de quelques dizaines de patients, le choix de la "
        "famille de modèle pèse davantage que le raffinement des variables. Le "
        "passage d'un réseau de 182 849 paramètres à une régression régularisée a "
        "produit, à variables strictement identiques, un écart de 0,94 point de "
        "corrélation. Les modèles à forte capacité, dont l'usage est aujourd'hui "
        "réflexe, sont ici contre-productifs.",
    )
    _para(
        doc,
        "Il faut enfin souligner que ce gain provient du modèle, et non d'une "
        "extraction de variables supérieure. Notre prétraitement demeure plus "
        "sommaire que celui de l'étude de référence.",
    )


def _limites(doc: Document) -> None:
    heading(doc, "7. Limites", 1)
    _para(
        doc,
        "Les réserves suivantes sont énoncées ici parce qu'elles conditionnent "
        "l'interprétation des chiffres du tableau précédent.",
    )
    bullet(
        doc,
        "Biais de sélection du modèle. La régression ridge a été retenue après "
        "comparaison de quatre familles de modèles sur les mêmes données ; son "
        "coefficient est donc optimiste. Un protocole de validation imbriqué est "
        "nécessaire pour l'estimer sans biais.",
    )
    bullet(
        doc,
        "Nombre de permutations insuffisant. Le contrôle repose sur "
        "50 permutations ; la queue de la distribution nulle atteignant +0,45, un "
        "millier de permutations serait requis pour une valeur de p stable.",
    )
    bullet(
        doc,
        "Protocole 2 non strictement comparable. Notre cohorte compte 88 patients "
        "contre 73 dans l'étude de référence. Les 15 sujets exclus par les auteurs "
        "ne sont pas identifiables à partir des métadonnées publiées ; nous avons "
        "choisi de conserver l'ensemble des patients disponibles plutôt que d'en "
        "écarter jusqu'à faire coïncider les effectifs.",
    )
    bullet(
        doc,
        "Prétraitement incomplet. L'analyse en composantes indépendantes, le rejet "
        "des segments artefactés et l'interpolation des électrodes défectueuses, "
        "présents dans l'étude de référence, ne sont pas implémentés.",
    )
    bullet(
        doc,
        "Absence d'échantillon de test indépendant. Toutes les valeurs rapportées "
        "proviennent d'une validation croisée sur la même cohorte. Une validation "
        "sur une cohorte externe reste nécessaire.",
    )


def _conclusion(doc: Document) -> None:
    heading(doc, "8. Conclusion", 1)
    _para(
        doc,
        "Ce travail établit qu'un signal EEG prédictif de la réponse à la rTMS est "
        "présent dans la cohorte TDBRAIN, et qu'il survit au retrait du principal "
        "facteur de confusion clinique. Il montre également que l'échec initial de "
        "notre approche provenait du choix du modèle et non des données — un "
        "diagnostic qui n'aurait pas été posé sans les contrôles décrits en "
        "section 4.",
    )
    _para(
        doc,
        "Les prolongements prioritaires sont, dans l'ordre : la levée du biais de "
        "sélection par validation imbriquée, l'augmentation du nombre de "
        "permutations, puis l'implémentation du prétraitement complet de l'étude "
        "de référence, seule voie susceptible d'améliorer encore les variables "
        "elles-mêmes.",
    )

    heading(doc, "Références", 1)
    _para(
        doc,
        "Arteaga et al. Prediction of rTMS treatment response in major depressive "
        "disorder from resting-state EEG. PMC12981298.",
    )
    _para(
        doc,
        "van Dijk H. et al. The Two Decades Brainclinics Research Archive for "
        "Insights in Neurophysiology (TDBRAIN) database. Scientific Data, 2022.",
    )
    doc.add_paragraph()
    _para(
        doc,
        "Toutes les valeurs citées sont reproductibles à partir du dépôt du "
        "projet : data/models/article_comparison.json, data/models/r_stability.json, "
        "et la procédure de validation croisée décrite en section 4.",
        italic=True,
    )


def build() -> Path:
    doc = Document()
    title_block(
        doc,
        "Améliorations méthodologiques",
        "Prédiction de la réponse au rTMS à partir d'un EEG de repos unique",
        "PFE 2026 — NeuroRéponse — comparaison à l'étude de référence PMC12981298",
    )
    doc.add_paragraph()
    _para(
        doc,
        "Note méthodologique destinée à un lecteur non informaticien. Elle expose "
        "ce qui a été modifié par rapport à l'étude de référence, pourquoi, et ce "
        "que les résultats obtenus permettent — ou ne permettent pas — de conclure.",
        italic=True,
    )
    doc.add_page_break()

    _resume(doc)
    _contexte(doc)
    _reference(doc)
    doc.add_page_break()
    _probleme(doc)
    doc.add_page_break()
    _ameliorations(doc)
    doc.add_page_break()
    _resultats(doc)
    _discussion(doc)
    doc.add_page_break()
    _limites(doc)
    _conclusion(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(f"écrit : {build()}")
