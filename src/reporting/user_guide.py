"""Build the user-facing DOCX guide.

    python -m src.reporting.capture_screens     # first: produce docs/screenshots/*.png
    python -m src.reporting.user_guide          # then: assemble the document

Covers every screen of the application, for both cohorts. Installation is
deliberately *not* repeated here — it lives in ``install_guide.py`` — so the two
documents cannot drift out of step with each other.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .docx_helpers import (
    bullet,
    code,
    figure,
    heading,
    note,
    step,
    table,
    title_block,
    warning,
)

SHOTS = Path("docs/screenshots")
OUTPUT = Path("docs/Guide_Utilisateur.docx")


def build() -> Path:
    doc = Document()

    title_block(
        doc,
        "Guide d'utilisation",
        "NeuroRéponse — application rTMS + LSTM (EEG / ECG)",
        "PFE 2026 — prise en main écran par écran",
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Ce guide présente l'application écran par écran : à quoi sert chaque page, "
        "comment l'utiliser, et comment lire ce qu'elle affiche. Il s'adresse à un "
        "utilisateur — clinicien, chercheur ou examinateur — qui découvre "
        "l'application."
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "L'installation n'est pas traitée ici : voir le « Guide d'installation ». "
        "Ce guide suppose l'application lancée et au moins une cohorte chargée."
    )
    r.italic = True

    doc.add_page_break()
    _principes(doc)
    doc.add_page_break()
    _accueil(doc)
    _patients(doc)
    doc.add_page_break()
    _sessions(doc)
    doc.add_page_break()
    _training(doc)
    doc.add_page_break()
    _predictions(doc)
    doc.add_page_break()
    _suivi(doc)
    doc.add_page_break()
    _boucle(doc)
    doc.add_page_break()
    _comparaison(doc)
    doc.add_page_break()
    _tdbrain(doc)
    _parcours(doc)
    _problemes(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


# --------------------------------------------------------------------------- #


def _principes(doc: Document) -> None:
    heading(doc, "1. Principes à connaître avant de commencer")

    heading(doc, "1.1 Trois cohortes, jamais mélangées", level=2)
    doc.add_paragraph(
        "La barre latérale porte deux sélecteurs, présents sur toutes les pages. "
        "« Cohorte » choisit les patients, « Jeu de variables » choisit le modèle "
        "appliqué à ces patients. Ensemble ils commandent l'intégralité de "
        "l'application : chaque cohorte a sa propre base SQLite, chaque modèle son "
        "propre fichier de poids. Un patient simulé et un patient réel ne peuvent "
        "donc jamais se retrouver dans la même table ni dans le même entraînement."
    )
    table(doc, ["", "Simulé — séquentiel", "Simulé — apparié", "TDBRAIN (réel)"], [
        ["Origine", "Générateur du projet", "Générateur calibré sur TDBRAIN",
         "Cohorte publique Brainclinics"],
        ["Patients", "100", "132", "132 (MDD traités par rTMS)"],
        ["Axe temporel", "10 séances de traitement", "8 époques d'un enregistrement",
         "8 époques d'un enregistrement"],
        ["Signal", "1 canal EEG", "26 canaux EEG + ECG", "26 canaux EEG + ECG"],
        ["Modèles", "1 (8 features)", "2 : clinique (4) ou multimodal (139)",
         "2 : clinique (4) ou multimodal (139)"],
        ["Base", "recherche.sqlite3", "recherche_sim_matched.sqlite3",
         "recherche_tdbrain.sqlite3"],
    ])
    note(
        doc,
        "Les deux cohortes simulées ne sont pas la même. La séquentielle est la "
        "seule où le modèle accumule l'information au fil des séances. L'appariée "
        "reproduit la structure de TDBRAIN et sert de contrôle : elle est générée "
        "sans effet neurophysiologique injecté.",
    )

    heading(doc, "1.2 « Séance » ne veut pas dire la même chose", level=2)
    warning(
        doc,
        "C'est le point le plus important pour lire correctement l'application. "
        "Dans la cohorte simulée, une séance est une vraie séance de traitement : "
        "il y a une évolution au fil du temps. Dans TDBRAIN, il n'existe qu'un seul "
        "enregistrement de repos réalisé AVANT le traitement ; ce qui est appelé "
        "« époque » est une fenêtre de ce même enregistrement. Il n'y a donc aucune "
        "trajectoire clinique à observer dans TDBRAIN.",
    )
    doc.add_paragraph(
        "L'application signale systématiquement cette différence à l'écran : les "
        "libellés passent de « séance » à « époque », et des encadrés rappellent que "
        "les paramètres rTMS et les scores BDI-II sont identiques d'une époque à "
        "l'autre."
    )

    heading(doc, "1.3 Le score utilisé", level=2)
    doc.add_paragraph(
        "La réponse au traitement est évaluée avec le BDI-II (échelle de dépression). "
        "Un patient est dit « répondeur » lorsque son score diminue d'au moins 50 % "
        "entre le début et la fin du traitement. C'est ce seuil de 50 % que le modèle "
        "apprend à prédire, et qui apparaît en pointillés sur les graphiques."
    )


def _accueil(doc: Document) -> None:
    heading(doc, "2. Page d'accueil")
    doc.add_paragraph(
        "Point d'entrée de l'application. Elle rappelle l'objet du projet, affiche "
        "l'état de la base pour la source sélectionnée, et permet d'initialiser la "
        "cohorte simulée en un clic si la base est vide."
    )
    figure(doc, SHOTS / "01_home.png", "Page d'accueil, cohorte simulée")
    bullet(doc, "Les sélecteurs « Cohorte » et « Jeu de variables » sont en haut de "
                "la barre latérale ; ils restent actifs quand on change de page.")
    bullet(doc, "La quatrième métrique rappelle le modèle actif : « Clinique seul (4) » "
                "ou « Multimodal — clinique + EEG + ECG (139) ».")
    bullet(doc, "Le bouton « Initialiser cette cohorte » génère le jeu de données s'il "
                "manque, puis l'insère en base.")
    note(
        doc,
        "Ce bouton ne concerne que la cohorte simulée. La cohorte réelle se charge "
        "en ligne de commande (voir le Guide d'installation, étape 8).",
    )


def _patients(doc: Document) -> None:
    heading(doc, "3. Page Patients")
    doc.add_paragraph(
        "Liste l'ensemble des patients de la cohorte active et permet de consulter "
        "le détail de chacun : identité, diagnostic, nombre de séances et historique "
        "clinique."
    )
    figure(doc, SHOTS / "02_patients.png", "Liste et détail d'un patient (cohorte simulée)")

    heading(doc, "Ce que l'on peut faire", level=2)
    step(doc, 1, "Consulter le tableau récapitulatif : identifiant, nom, âge, "
                 "diagnostic, nombre de séances et dernier score enregistré.")
    step(doc, 2, "Sélectionner un patient dans la liste déroulante « Détail d'un "
                 "patient » pour afficher sa fiche et son historique clinique.")
    step(doc, 3, "Ajouter une entrée au dossier (note + score) via l'encadré "
                 "« Ajouter une entrée au dossier ».")
    step(doc, 4, "Créer un nouveau patient via « Créer un nouveau patient ».")
    note(
        doc,
        "Un patient créé manuellement n'a aucun signal EEG : il apparaîtra dans la "
        "liste mais ne pourra pas servir à une prédiction tant qu'aucune séance avec "
        "signal ne lui a été ajoutée (voir la page Boucle clinique).",
    )


def _sessions(doc: Document) -> None:
    heading(doc, "4. Page Sessions")
    doc.add_paragraph(
        "Détaille une séance (ou une époque) d'un patient : son statut, les scores "
        "avant/après, les paramètres de stimulation appliqués et les signaux "
        "enregistrés."
    )
    figure(doc, SHOTS / "03_sessions.png", "Détail d'une séance (cohorte simulée)")

    heading(doc, "4.1 Paramètres rTMS", level=2)
    doc.add_paragraph(
        "Le bloc « Paramètres rTMS » indique la fréquence, l'intensité, la durée et "
        "le nombre de trains, la localisation et le nombre total d'impulsions."
    )
    warning(
        doc,
        "Sur la cohorte TDBRAIN, plusieurs de ces valeurs s'affichent "
        "« — (non publié) ». Ce n'est pas un défaut de l'application : la base "
        "source publie le protocole (fréquence et site de stimulation) mais pas la "
        "dose administrée patient par patient. Ces champs sont volontairement laissés "
        "vides plutôt que remplis avec des valeurs plausibles mais inventées.",
    )

    heading(doc, "4.2 Signaux enregistrés", level=2)
    doc.add_paragraph(
        "La liste déroulante « Canal » donne accès à chaque signal de la séance. "
        "Sur la cohorte simulée il y a un canal EEG. Sur TDBRAIN, chaque époque "
        "contient 27 signaux : les 26 canaux du montage EEG et la dérivation ECG."
    )
    figure(doc, SHOTS / "td_03_sessions.png",
           "Séance TDBRAIN : 27 signaux, paramètres de dose non publiés")

    heading(doc, "4.3 Le signal ECG (tachogramme RR)", level=2)
    doc.add_paragraph(
        "En sélectionnant le canal « Erbs (ECG) », l'application affiche le "
        "tachogramme : la durée entre battements cardiaques successifs, en "
        "millisecondes. L'axe horizontal est le numéro de battement, et non le temps, "
        "car un rythme cardiaque n'est pas échantillonné régulièrement."
    )
    figure(doc, SHOTS / "td_03b_sessions_ecg.png",
           "Tachogramme RR et métriques HRV extraites de la dérivation ECG")
    doc.add_paragraph(
        "L'encadré « Features extraits » donne les cinq métriques de variabilité "
        "cardiaque calculées : fréquence cardiaque moyenne, SDNN, RMSSD, pNN50 et "
        "rapport LF/HF. Ce sont exactement les cinq valeurs que le modèle utilise "
        "lorsque la modalité ECG est activée."
    )
    note(
        doc,
        "La variabilité cardiaque est mesurée sur l'enregistrement complet (environ "
        "2 minutes) puis reportée à l'identique sur chaque époque : c'est une "
        "caractéristique du patient, pas une évolution dans le temps. Une époque de "
        "8 secondes ne contient qu'une dizaine de battements, ce qui serait "
        "insuffisant pour un calcul fiable.",
    )


def _training(doc: Document) -> None:
    heading(doc, "5. Page Training (entraînement du modèle)")
    doc.add_paragraph(
        "C'est ici que le modèle LSTM est évalué puis entraîné. La page affiche "
        "d'abord la composition de la cohorte et le nombre de features, puis propose "
        "les hyperparamètres et deux onglets."
    )
    figure(doc, SHOTS / "04_training_form.png", "Hyperparamètres et onglets (cohorte simulée)")

    heading(doc, "5.1 Onglet « Validation croisée »", level=2)
    doc.add_paragraph(
        "Mesure la performance du modèle sans l'enregistrer. La validation est "
        "toujours effectuée patient par patient (GroupKFold) : les séances d'un même "
        "patient ne sont jamais réparties entre l'entraînement et le test, ce qui "
        "évite de surestimer la performance."
    )
    figure(doc, SHOTS / "04_training.png", "Résultats de validation croisée (cohorte simulée)")
    doc.add_paragraph("Trois indicateurs sont affichés :")
    table(doc, ["Indicateur", "Lecture"], [
        ["Accuracy", "Proportion de patients correctement classés"],
        ["AUC", "Capacité à séparer répondeurs et non-répondeurs ; 50 % = hasard"],
        ["F1", "Compromis entre précision et rappel sur la classe répondeur"],
    ])

    heading(doc, "5.2 Onglet « Modèle final »", level=2)
    doc.add_paragraph(
        "Entraîne le modèle sur l'ensemble des patients et l'enregistre sur disque. "
        "C'est cette étape qui rend les pages Predictions, Suivi et Boucle clinique "
        "utilisables."
    )

    heading(doc, "5.3 Options propres à la cohorte réelle", level=2)
    doc.add_paragraph(
        "Sous TDBRAIN, deux cases apparaissent dans la barre latérale : la "
        "normalisation z-score intra-patient, et l'ajout de la modalité ECG (HRV). "
        "Le compteur « Features » passe de 130 à 135 lorsque l'ECG est activé."
    )
    figure(doc, SHOTS / "td_04_training.png",
           "Training sur TDBRAIN : 132 patients, 135 features (EEG + HRV)")


def _predictions(doc: Document) -> None:
    heading(doc, "6. Page Predictions")
    doc.add_paragraph(
        "Applique le modèle entraîné à un patient et affiche la probabilité qu'il "
        "réponde au traitement."
    )
    figure(doc, SHOTS / "05_predictions.png", "Prédiction pour un patient (cohorte simulée)")

    heading(doc, "Lecture de l'écran", level=2)
    bullet(doc, "« Probabilité de réponse » : sortie du modèle, entre 0 et 100 %.")
    bullet(doc, "« Classe prédite » : répondeur au-dessus de 50 %, non-répondeur en dessous.")
    bullet(doc, "« Indice thérapeutique (TRI) » : la probabilité recalculée après "
                "chaque séance, à mesure que le modèle accumule de l'information.")
    bullet(doc, "« Comparaison au score clinique observé » : confronte la prédiction "
                "à la réduction BDI-II réellement mesurée (concordance vraie ou fausse).")
    doc.add_paragraph(
        "Les deux boutons du bas permettent d'enregistrer la prédiction en base et "
        "de télécharger un rapport PDF."
    )
    figure(doc, SHOTS / "td_05_predictions.png", "Prédiction sur un patient réel TDBRAIN")
    note(
        doc,
        "Sur TDBRAIN, la courbe TRI porte sur des époques d'un même enregistrement : "
        "elle montre l'accumulation de preuve du modèle sur le signal, et non une "
        "amélioration clinique au fil du traitement. L'application le rappelle sous "
        "le graphique.",
    )


def _suivi(doc: Document) -> None:
    heading(doc, "7. Page Suivi")
    doc.add_paragraph(
        "Contrairement à la page Predictions, qui produit un résultat unique, la page "
        "Suivi synthétise l'ensemble des séances d'un patient : trajectoire clinique, "
        "trajectoire du modèle, et cohérence entre les deux."
    )
    figure(doc, SHOTS / "06_suivi.png", "Suivi d'un patient sur ses 10 séances (cohorte simulée)")

    heading(doc, "7.1 Trajectoire clinique", level=2)
    doc.add_paragraph(
        "Le premier graphique trace le score BDI-II séance après séance. Deux repères "
        "l'accompagnent : le score de départ (pointillés gris) et le seuil de réponse "
        "à −50 % (pointillés verts). Une pente est estimée et qualifiée "
        "d'amélioration, stable ou aggravation."
    )

    heading(doc, "7.2 Trajectoire du modèle", level=2)
    doc.add_paragraph(
        "Le second graphique trace l'indice TRI. Trois indicateurs le résument : TRI "
        "final, TRI moyen et écart-type entre séances. Un écart-type élevé signale une "
        "prédiction instable, à interpréter avec prudence."
    )

    heading(doc, "7.3 Retour de synthèse", level=2)
    doc.add_paragraph(
        "En bas de page, l'application produit un commentaire en langage clair : "
        "tendance observée, réduction totale, stabilité du modèle et concordance avec "
        "l'issue réelle. Un tableau détaille ensuite chaque séance."
    )
    figure(doc, SHOTS / "td_06_suivi.png", "Suivi d'un patient TDBRAIN")
    warning(
        doc,
        "Sur TDBRAIN, un bandeau indique « Pas de trajectoire clinique ». C'est "
        "attendu : les époques portent toutes le même score BDI-II, puisqu'il n'existe "
        "qu'un seul enregistrement avant traitement. L'application refuse alors de "
        "calculer une tendance, et présente l'écart-type du TRI comme une mesure de "
        "cohérence entre fenêtres, non comme un progrès.",
    )


def _boucle(doc: Document) -> None:
    heading(doc, "8. Page Boucle clinique")
    doc.add_paragraph(
        "Cette page met en œuvre le scénario d'usage visé : enregistrer une séance, "
        "obtenir une prédiction, ajuster le stimulateur, puis recommencer — jusqu'à "
        "obtenir un résultat satisfaisant."
    )
    doc.add_paragraph("Le cycle se déroule ainsi :")
    step(doc, 1, "Enregistrer la séance : paramètres appliqués sur la machine, "
                 "signal EEG et score clinique.")
    step(doc, 2, "L'application recalcule la prédiction et l'ajoute au journal.")
    step(doc, 3, "Le réglage du stimulateur est modifié sur l'appareil — cette "
                 "opération se fait hors application.")
    step(doc, 4, "La séance suivante est enregistrée, et ainsi de suite.")
    figure(doc, SHOTS / "07_boucle.png", "Boucle clinique, cohorte simulée")

    heading(doc, "8.1 Enregistrer une séance", level=2)
    doc.add_paragraph(
        "Le formulaire est pré-rempli avec les paramètres de la séance précédente : "
        "il suffit de modifier ce qui a changé sur le stimulateur. Le signal peut être "
        "téléversé (fichier) ou généré à des fins de démonstration lorsqu'aucun "
        "matériel n'est connecté."
    )
    note(
        doc,
        "L'option « Générer (démonstration) » produit un signal synthétique. Elle "
        "sert uniquement à dérouler le cycle sans amplificateur ; elle ne constitue "
        "en aucun cas une donnée clinique.",
    )

    heading(doc, "8.2 Deux modes selon la cohorte", level=2)
    doc.add_paragraph(
        "Les deux cohortes acceptent la boucle, mais accumulent l'information "
        "différemment, car les deux modèles n'ont pas été entraînés sur le même axe."
    )
    table(doc, ["", "Données simulées", "TDBRAIN"], [
        ["Mode", "Séquentiel", "Instantané"],
        ["Entrée du modèle", "La suite des séances", "L'enregistrement de la séance seule"],
        ["Accumulation", "À l'intérieur du LSTM", "Au niveau clinique (tendance)"],
        ["Fichier attendu", "CSV ou NPY, un canal", "BDF ou CSV, 26 canaux + ECG"],
    ])
    doc.add_paragraph(
        "En mode instantané, chaque séance exige un nouvel enregistrement de repos "
        "d'au moins 64 secondes. Le fichier subit exactement le même prétraitement "
        "que les données d'entraînement, ce qui garantit que le modèle reçoit des "
        "valeurs comparables à celles qu'il a apprises."
    )
    figure(doc, SHOTS / "td_07_boucle.png", "Boucle clinique en mode instantané (TDBRAIN)")

    heading(doc, "8.3 Journal et recommandation", level=2)
    doc.add_paragraph(
        "Le graphique retrace l'évolution de la prédiction. Le journal associe à "
        "chaque séance les paramètres utilisés, la probabilité obtenue, l'écart avec "
        "la séance précédente et les réglages modifiés — ce qui permet de voir quel "
        "ajustement a produit quel effet."
    )
    doc.add_paragraph(
        "La recommandation indique la direction observée : conserver les paramètres, "
        "revenir aux précédents, ou constater un plateau."
    )
    warning(
        doc,
        "La recommandation ne propose jamais un réglage chiffré. Aucune donnée de ce "
        "projet ne relie une intensité de stimulation à la réponse — la base TDBRAIN "
        "ne publie même pas la dose administrée. Le choix des paramètres reste une "
        "décision clinique.",
    )


def _comparaison(doc: Document) -> None:
    heading(doc, "9. Page Comparaison")
    doc.add_paragraph(
        "Cette page met les quatre modèles côte à côte : deux cohortes (simulée "
        "appariée, TDBRAIN réelle) croisées avec deux jeux de variables (clinique "
        "seul, clinique + EEG + ECG). C'est le résultat central du projet, et la "
        "page ne recalcule rien : elle lit le fichier produit par l'entraînement, "
        "donc les chiffres affichés sont exactement ceux qui ont été mesurés."
    )
    figure(doc, SHOTS / "08_comparaison.png", "Les quatre modèles côte à côte")

    heading(doc, "9.1 Comment lire les deux graphiques", level=2)
    bullet(doc, "Le premier montre l'AUC de chaque modèle avec l'écart-type entre "
                "les plis. La ligne pointillée est le hasard (0,5).")
    bullet(doc, "Le second compare l'exactitude au taux de la classe majoritaire. "
                "C'est le graphique qui compte : une exactitude égale à ce taux "
                "signifie que le modèle prédit toujours la même classe.")

    heading(doc, "9.2 Ce que la page conclut", level=2)
    bullet(doc, "Aucune variante ne dépasse son taux de base. Le modèle répond "
                "« répondeur » pour presque tous les patients.")
    bullet(doc, "Le modèle clinique (4 variables) fait aussi bien, voire mieux, que "
                "le multimodal (139) : ajouter des colonnes non informatives dilue "
                "les quelques-unes qui le sont, sur seulement 132 patients.")
    bullet(doc, "Les écarts entre variantes sont de l'ordre de l'écart-type entre "
                "plis : les classer par AUC reviendrait à commenter du bruit.")
    note(
        doc,
        "Un avertissement s'affiche automatiquement lorsque aucune variante ne "
        "dépasse son taux de base. Il n'est pas décoratif : c'est le résultat, et "
        "il est négatif. La démonstration porte sur la chaîne de traitement "
        "complète, pas sur une performance clinique.",
    )


def _tdbrain(doc: Document) -> None:
    heading(doc, "10. Particularités de la cohorte réelle")
    figure(doc, SHOTS / "td_01_home.png", "Accueil avec la cohorte TDBRAIN sélectionnée")

    heading(doc, "10.1 Ce qui est réel, ce qui ne l'est pas", level=2)
    table(doc, ["Élément", "Statut"], [
        ["EEG 26 canaux, dérivation ECG", "Réel (mesuré)"],
        ["Scores BDI-II avant / après", "Réels"],
        ["Étiquette répondeur", "Réelle (réduction ≥ 50 %)"],
        ["Âge, sexe", "Réels"],
        ["Protocole rTMS (fréquence, site)", "Réel"],
        ["Intensité, trains, intervalle", "Non publiés — affichés « — »"],
        ["« Époques »", "Fenêtres d'un enregistrement unique, pas des séances"],
    ])

    heading(doc, "10.2 Le résultat obtenu, et comment le présenter", level=2)
    doc.add_paragraph(
        "Sur la cohorte réelle, la prédiction de la réponse au rTMS se situe au "
        "niveau du hasard : l'AUC mesurée se situe entre 0,47 et 0,52 selon les "
        "variantes, et l'exactitude reste égale au taux de base de la cohorte. "
        "L'ajout de la modalité ECG ne change pas cette conclusion."
    )
    doc.add_paragraph(
        "Ce résultat n'est pas un défaut de l'application. Il s'explique par la nature "
        "des données : un unique enregistrement de repos réalisé avant traitement ne "
        "contient pas, à lui seul, l'information nécessaire pour anticiper la réponse. "
        "C'est un résultat négatif honnête, cohérent avec la littérature, et il doit "
        "être présenté comme tel plutôt que masqué."
    )
    note(
        doc,
        "À titre de comparaison, la même chaîne de traitement appliquée à une tâche "
        "différente — distinguer patients dépressifs et sujets sains — atteint une "
        "AUC nettement supérieure. Le problème vient donc bien de la tâche visée, non "
        "de la chaîne de traitement.",
    )


def _parcours(doc: Document) -> None:
    heading(doc, "11. Parcours conseillé pour une démonstration")
    step(doc, 1, "Accueil : vérifier la source sélectionnée et l'état de la base.")
    step(doc, 2, "Patients : ouvrir un patient, montrer son historique clinique.")
    step(doc, 3, "Sessions : montrer les paramètres rTMS, un canal EEG, puis — sur "
                 "TDBRAIN — le tachogramme RR et les métriques HRV.")
    step(doc, 4, "Training : lancer une validation croisée et commenter les trois "
                 "indicateurs.")
    step(doc, 5, "Predictions : produire une prédiction et la confronter à l'issue "
                 "réelle du patient.")
    step(doc, 6, "Suivi : montrer la synthèse sur l'ensemble des séances.")
    step(doc, 7, "Boucle clinique : enregistrer une séance et montrer la prédiction "
                 "se mettre à jour.")
    doc.add_paragraph(
        "Basculer entre les deux cohortes au moment de l'étape 3 met en évidence la "
        "différence de nature des données, qui est le point clé du projet."
    )


def _problemes(doc: Document) -> None:
    heading(doc, "12. Messages fréquents")
    table(doc, ["Message affiché", "Signification", "Que faire"], [
        ["Aucun patient en base",
         "La cohorte n'a pas été chargée",
         "Accueil : « Initialiser avec les données simulées »"],
        ["Aucun modèle entraîné pour cette source",
         "Aucun modèle n'existe encore",
         "Page Training, onglet « Modèle final »"],
        ["Les données en base ne correspondent pas au modèle",
         "Base et modèle TDBRAIN désynchronisés",
         "Recharger la cohorte et réentraîner (Guide d'installation, étape 8)"],
        ["Pas de trajectoire clinique",
         "Comportement normal sur TDBRAIN",
         "Aucune action : il n'existe qu'un enregistrement avant traitement"],
        ["— (non publié)",
         "Donnée absente de la base source",
         "Aucune action : la valeur n'est volontairement pas inventée"],
        ["appartient à la cohorte de recherche TDBRAIN",
         "Ce patient n'a pas de séance de traitement",
         "Créer un nouveau patient pour démarrer une boucle clinique"],
        ["Enregistrement trop court",
         "Le fichier fourni dure moins de 64 secondes",
         "Fournir un enregistrement plus long"],
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Pour toute question d'installation ou de mise à jour des données, se "
        "reporter au « Guide d'installation »."
    )
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph()
    heading(doc, "Régénérer ce document", level=2)
    code(doc, [
        "streamlit run src/app/main.py --server.headless true --server.port 8765",
        "python -m src.reporting.capture_screens",
        "python -m src.reporting.user_guide",
    ])


if __name__ == "__main__":
    print(f"Wrote {build()}")
