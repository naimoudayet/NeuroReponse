"""Shared python-docx building blocks for the generated guides.

Both ``install_guide.py`` and ``user_guide.py`` produce Word documents in the same
house style. Keeping the primitives here means a change to, say, the code-block
font applies to both instead of drifting between them.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(0x1E, 0x3A, 0x8A)
ORANGE = RGBColor(0xB4, 0x53, 0x09)
GREY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0x99, 0x1B, 0x1B)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE


def title_block(doc: Document, title: str, subtitle: str, tag: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(tag)
    r3.italic = True
    r3.font.size = Pt(11)


def step(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Étape {n}. ")
    r.bold = True
    p.add_run(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Note : ")
    r.bold = True
    r.font.color.rgb = ORANGE
    p.add_run(text)


def warning(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Attention : ")
    r.bold = True
    r.font.color.rgb = RED
    p.add_run(text)


def code(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = GREY


def figure(doc: Document, image: Path, caption: str, width_cm: float = 15.0) -> None:
    """Embed a screenshot, or a visible placeholder when it has not been captured.

    A missing file is flagged rather than skipped: silently omitting a figure makes
    an incomplete guide look complete.
    """
    if not image.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[capture manquante : {image.name} — lancer capture_screens.py]")
        r.italic = True
        r.font.color.rgb = RED
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure : {caption}")
    r.italic = True
    r.font.size = Pt(9)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
    doc.add_paragraph()
