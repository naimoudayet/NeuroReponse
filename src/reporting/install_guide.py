"""Build the installation / deployment DOCX guide.

    python -m src.reporting.install_guide

Documents how to bring the application up on a **fresh machine**. Everything it
claims is checked against the repository as it stands: the dependency list, the
CLI entry points that actually exist, and which artifacts are git-ignored and must
therefore be regenerated rather than copied.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from .docx_helpers import (
    bullet,
    code,
    heading,
    note,
    step,
    table,
    title_block,
    warning,
)

OUTPUT = Path("docs/Guide_Installation.docx")


def build() -> Path:
    doc = Document()

    title_block(
        doc,
        "Guide d'installation",
        "NeuroRéponse — application rTMS + LSTM (EEG / ECG)",
        "PFE 2026 — déploiement sur une nouvelle machine",
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Ce document décrit, pas à pas, l'installation complète de l'application "
        "sur une machine vierge : dépendances, initialisation de la base, "
        "entraînement du modèle et lancement de l'interface. Il couvre aussi la "
        "préparation optionnelle de la cohorte réelle TDBRAIN. Pour l'utilisation "
        "de l'application une fois installée, voir le « Guide d'utilisation »."
    )

    _apercu(doc)
    doc.add_page_break()
    _prerequis(doc)
    _etapes(doc)
    doc.add_page_break()
    _tdbrain(doc)
    _documentation(doc)
    doc.add_page_break()
    _artefacts(doc)
    _depannage(doc)
    _recapitulatif(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


# --------------------------------------------------------------------------- #


def _apercu(doc: Document) -> None:
    heading(doc, "1. Aperçu de l'installation")
    doc.add_paragraph(
        "L'installation se fait en deux temps. La cohorte simulée suffit pour "
        "démarrer et n'exige aucun téléchargement externe ; la cohorte réelle "
        "TDBRAIN est optionnelle et soumise à un accord d'utilisation."
    )
    table(doc, ["Étape", "Durée indicative", "Obligatoire ?"], [
        ["1 à 3 — Python, environnement, dépendances", "10–20 min", "Oui"],
        ["4 — Vérification (tests)", "2 min", "Recommandé"],
        ["5 — Initialiser la cohorte simulée", "< 1 min", "Oui"],
        ["6 — Entraîner le modèle simulé", "1–3 min", "Oui (sinon pas de prédiction)"],
        ["7 — Lancer l'application", "immédiat", "Oui"],
        ["8 — Cohorte réelle TDBRAIN", "≈ 14,5 Go + 10–30 min", "Non (optionnel)"],
        ["9 — Régénérer la documentation", "5 min", "Non (optionnel)"],
    ])


def _prerequis(doc: Document) -> None:
    heading(doc, "2. Prérequis")

    heading(doc, "2.1 Logiciels", level=2)
    bullet(doc, "Python 3.11 ou plus récent — développé et testé sur Python 3.14.")
    bullet(doc, "pip (fourni avec Python) et le module venv.")
    bullet(doc, "Git, si le projet est récupéré depuis un dépôt.")
    bullet(doc, "Un navigateur web récent : l'interface Streamlit s'ouvre dedans.")
    note(
        doc,
        "TensorFlow n'est pas utilisé : il ne fournit pas de paquet pour Python 3.14. "
        "Le modèle LSTM est écrit en PyTorch (CPU), l'architecture et les métriques "
        "sont identiques à celles décrites dans le document de conception.",
    )

    heading(doc, "2.2 Matériel", level=2)
    bullet(doc, "Environ 2 Go d'espace disque pour Python + PyTorch + dépendances.")
    bullet(doc, "Aucune carte graphique requise : tout tourne sur CPU.")
    bullet(doc, "4 Go de RAM suffisent pour la cohorte simulée.")
    bullet(
        doc,
        "Pour la cohorte réelle TDBRAIN : environ 15 Go supplémentaires pour "
        "l'archive, plus ≈ 250 Mo pour la base SQLite générée.",
    )

    heading(doc, "2.3 Systèmes d'exploitation", level=2)
    doc.add_paragraph(
        "Le projet est développé sous Windows 11 (PowerShell). Il fonctionne aussi "
        "sous Linux et macOS ; seules les commandes d'activation de l'environnement "
        "virtuel changent, elles sont indiquées à l'étape 2."
    )


def _etapes(doc: Document) -> None:
    heading(doc, "3. Installation pas à pas")

    heading(doc, "Étape 1 — Récupérer le projet", level=2)
    doc.add_paragraph("Depuis un dépôt Git :")
    code(doc, [
        "git clone <url-du-depot> NeuroReponse",
        "cd NeuroReponse",
    ])
    doc.add_paragraph(
        "Ou bien copier le dossier du projet, puis ouvrir une console dans ce dossier. "
        "Toutes les commandes de ce guide s'exécutent depuis la racine du projet "
        "(le dossier qui contient requirements.txt)."
    )

    heading(doc, "Étape 2 — Créer un environnement virtuel", level=2)
    doc.add_paragraph("Windows (PowerShell) :")
    code(doc, ["python -m venv .venv", ".\\.venv\\Scripts\\Activate.ps1"])
    doc.add_paragraph("Linux / macOS :")
    code(doc, ["python3 -m venv .venv", "source .venv/bin/activate"])
    note(
        doc,
        "Si PowerShell refuse d'exécuter le script d'activation, autoriser les "
        "scripts pour la session en cours : "
        "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass",
    )

    heading(doc, "Étape 3 — Installer les dépendances", level=2)
    step(doc, 1, "Mettre pip à jour, puis installer les dépendances du projet :")
    code(doc, [
        "python -m pip install --upgrade pip",
        "pip install -r requirements.txt",
    ])
    step(doc, 2, "Si PyTorch n'est pas disponible pour votre plateforme via la "
                 "commande précédente, l'installer depuis l'index CPU officiel :")
    code(doc, ["pip install torch --index-url https://download.pytorch.org/whl/cpu"])
    note(
        doc,
        "requirements.txt installe aussi mne (lecture des fichiers EEG BioSemi BDF), "
        "python-docx et playwright. Ces deux derniers ne servent qu'à régénérer la "
        "documentation (étape 9) ; ils n'empêchent pas l'application de fonctionner.",
    )

    heading(doc, "Étape 4 — Vérifier l'installation", level=2)
    doc.add_paragraph(
        "Cette étape n'est pas obligatoire mais confirme en deux minutes que "
        "l'environnement est sain avant d'aller plus loin."
    )
    code(doc, [
        "python -m pytest -q",
        "python -m ruff check src/ --select=E,F,I,UP --ignore=E501",
    ])
    doc.add_paragraph(
        "Résultat attendu : la totalité des tests passe et l'analyse statique ne "
        "signale rien. Les tests qui concernent la cohorte réelle utilisent un jeu "
        "de données synthétique et ne demandent aucun téléchargement."
    )

    heading(doc, "Étape 5 — Initialiser la cohorte simulée", level=2)
    doc.add_paragraph("Deux possibilités, au choix.")
    doc.add_paragraph("Option A — depuis l'application (la plus simple) :")
    bullet(doc, "Lancer l'application (étape 7), puis cliquer sur le bouton "
                "« Initialiser avec les données simulées » de la page d'accueil. "
                "Le jeu de données est généré s'il n'existe pas, puis inséré en base.")
    doc.add_paragraph("Option B — en ligne de commande :")
    code(doc, [
        "python -m src.data.simulator     # génère data/simulated/",
        "python -m src.data.seeder        # insère dans recherche.sqlite3",
        "",
        "# Cohorte simulée *appariée* sur TDBRAIN (132 patients, 26 canaux + ECG),",
        "# nécessaire aux deux modèles simulés de la comparaison 2x2 :",
        "python -m src.data.tdbrain_seeder --matched",
    ])
    note(
        doc,
        "Le dossier data/simulated/ et les fichiers *.sqlite3 sont exclus du dépôt "
        "Git : sur une machine neuve ils n'existent pas et doivent être régénérés "
        "par cette étape. C'est normal.",
    )

    heading(doc, "Étape 6 — Entraîner le modèle simulé", level=2)
    warning(
        doc,
        "Il n'existe pas de commande en ligne pour entraîner le modèle simulé : "
        "il s'entraîne depuis l'interface. Sans cette étape, la page Predictions "
        "affichera « aucun modèle entraîné » et la Boucle clinique sera bloquée.",
    )
    step(doc, 1, "Lancer l'application (étape 7) et ouvrir la page « Training ».")
    step(doc, 2, "Vérifier que la source « Données simulées » est sélectionnée "
                 "dans la barre latérale.")
    step(doc, 3, "Ouvrir l'onglet « Modèle final », puis lancer l'entraînement. "
                 "Le fichier data/models/lstm_v1.pt est écrit à la fin.")
    doc.add_paragraph(
        "L'onglet « Validation croisée » sert à mesurer la performance "
        "(GroupKFold patient-wise) ; il n'enregistre pas de modèle."
    )

    heading(doc, "Étape 7 — Lancer l'application", level=2)
    code(doc, ["streamlit run src/app/main.py"])
    doc.add_paragraph(
        "Le navigateur s'ouvre sur http://localhost:8501. Pour choisir un autre "
        "port ou lancer sans ouvrir de navigateur :"
    )
    code(doc, [
        "streamlit run src/app/main.py --server.port 8600 --server.headless true",
    ])
    note(
        doc,
        "Si la commande « streamlit » est introuvable alors que l'installation a "
        "réussi, utiliser python -m streamlit run src/app/main.py — cela contourne "
        "un PATH mal configuré.",
    )


def _tdbrain(doc: Document) -> None:
    heading(doc, "4. Étape 8 (optionnelle) — Cohorte réelle TDBRAIN")
    doc.add_paragraph(
        "L'application fonctionne pleinement avec la seule cohorte simulée. Cette "
        "étape ajoute la cohorte réelle TDBRAIN : 132 patients dépressifs traités "
        "par rTMS, avec EEG 26 canaux et dérivation ECG."
    )

    warning(
        doc,
        "Les données TDBRAIN sont soumises à un accord d'utilisation (Data Use "
        "Agreement). Elles ne doivent jamais être ajoutées au dépôt Git, ni "
        "redistribuées. Le fichier .gitignore les exclut déjà.",
    )

    heading(doc, "8.1 Obtenir les données", level=2)
    step(doc, 1, "Se rendre sur https://brainclinics.com/resources/ et ouvrir le "
                 "jeu de données TDBRAIN.")
    step(doc, 2, "Créer un compte (ORCID) et accepter l'accord d'utilisation.")
    step(doc, 3, "Télécharger « TDBRAIN Dataset V3.1 » (archive chiffrée, ≈ 14,5 Go), "
                 "puis la décompresser.")
    note(
        doc,
        "Les paquets publics « Treatment / Diagnostic Prediction » sont anonymisés "
        "quant aux étiquettes de réponse : seules les données complètes V3.1 "
        "contiennent les vraies étiquettes répondeur / non-répondeur.",
    )

    heading(doc, "8.2 Placer les données", level=2)
    doc.add_paragraph(
        "Placer le dossier décompressé sous data/tdbrain/. Le dossier attendu est "
        "celui qui contient participants.tsv et les dossiers sub-XXXXXXXX/ :"
    )
    code(doc, [
        "data/tdbrain/TDBRAIN_Dataset_V3_1_Encr/TDBRAIN_Dataset_V3_1/",
        "    participants.tsv",
        "    dataset_description.json",
        "    sub-87999321/ses-1/eeg/sub-87999321_ses-1_task-restEO_eeg.bdf",
        "    ...",
    ])

    heading(doc, "8.3 Charger et entraîner en une passe", level=2)
    doc.add_paragraph(
        "La commande suivante lit un fichier BDF par patient, insère la cohorte en "
        "base, évalue les variantes (EEG seul / EEG + ECG, avec et sans "
        "normalisation) et enregistre le meilleur modèle avec son contrat de "
        "features. Comptez plusieurs minutes : la lecture des BDF est le poste "
        "le plus long."
    )
    code(doc, [
        "python -m src.models.train_tdbrain `",
        '    --root "data/tdbrain/TDBRAIN_Dataset_V3_1_Encr/TDBRAIN_Dataset_V3_1" `',
        "    --seed-db recherche_tdbrain.sqlite3",
    ])
    doc.add_paragraph("Options utiles :")
    table(doc, ["Option", "Effet"], [
        ["--modalities eeg", "N'utiliser que l'EEG (130 features)"],
        ["--modalities eeg+ecg", "Forcer EEG + HRV (135 features)"],
        ["--no-ecg", "Ne pas lire la dérivation ECG (chargement plus rapide)"],
        ["--zscore on|off|auto", "Normalisation z-score intra-patient"],
        ["--n-splits 5", "Nombre de plis de validation croisée"],
    ])
    note(
        doc,
        "Chaque modèle TDBRAIN est enregistré avec un fichier compagnon JSON "
        "(le « contrat de features ») décrivant canaux, fréquence, époques et "
        "modalités. L'application refuse de prédire si les données en base ne "
        "correspondent pas à ce contrat — ne pas supprimer ce fichier.",
    )
    doc.add_paragraph(
        "Les trois cohortes utilisent des bases SQLite distinctes "
        "(recherche.sqlite3, recherche_sim_matched.sqlite3 et "
        "recherche_tdbrain.sqlite3) et des modèles distincts : patients simulés et "
        "réels ne se mélangent jamais."
    )
    doc.add_paragraph(
        "Pour entraîner les quatre modèles de la comparaison en une commande "
        "(deux cohortes × deux jeux de variables) :"
    )
    code(doc, [
        "python -m src.models.train_all `",
        '    --root "data/tdbrain/TDBRAIN_Dataset_V3_1_Encr/TDBRAIN_Dataset_V3_1"',
        "",
        "python -m src.models.train_all --sim-only   # sans la cohorte réelle",
    ])


def _documentation(doc: Document) -> None:
    heading(doc, "5. Étape 9 (optionnelle) — Régénérer la documentation")
    doc.add_paragraph(
        "Les deux guides (installation et utilisation) sont générés par script. "
        "Les captures d'écran du guide d'utilisation sont prises automatiquement "
        "sur l'application en fonctionnement."
    )
    step(doc, 1, "Installer le navigateur utilisé par Playwright (une seule fois) :")
    code(doc, ["python -m playwright install chromium"])
    step(doc, 2, "Lancer l'application sur le port 8765 (celui qu'attend le script) :")
    code(doc, [
        "streamlit run src/app/main.py --server.headless true --server.port 8765",
    ])
    step(doc, 3, "Dans une seconde console, capturer les écrans puis régénérer "
                 "les documents :")
    code(doc, [
        "python -m src.reporting.capture_screens",
        "python -m src.reporting.user_guide",
        "python -m src.reporting.install_guide",
    ])
    note(
        doc,
        "Les captures sont prises pour les deux cohortes. Si la cohorte réelle "
        "n'est pas installée, les captures TDBRAIN manqueront et le guide affichera "
        "un repère « capture manquante » à leur emplacement.",
    )


def _artefacts(doc: Document) -> None:
    heading(doc, "6. Ce qui est versionné, ce qui est régénéré")
    doc.add_paragraph(
        "Sur une machine neuve, plusieurs fichiers sont absents volontairement. "
        "Les recréer fait partie de l'installation ; ce n'est pas une anomalie."
    )
    table(doc, ["Élément", "Versionné ?", "Comment l'obtenir"], [
        ["Code source, tests, notebooks", "Oui", "Fourni avec le projet"],
        ["data/simulated/*.npz, *.csv", "Non", "python -m src.data.simulator"],
        ["recherche.sqlite3", "Non", "Bouton d'accueil ou src.data.seeder"],
        ["recherche_sim_matched.sqlite3", "Non", "tdbrain_seeder --matched"],
        ["recherche_tdbrain.sqlite3", "Non", "train_tdbrain --seed-db"],
        ["data/models/lstm_v1.pt", "Non", "Page Training, onglet « Modèle final »"],
        ["data/models/sim_*.pt, tdbrain_*.pt (+ .json)", "Non", "train_all"],
        ["data/tdbrain/ (données réelles)", "Non — interdit", "Téléchargement sous accord"],
        ["docs/screenshots/*.png", "Non", "python -m src.reporting.capture_screens"],
    ])


def _depannage(doc: Document) -> None:
    heading(doc, "7. Dépannage")
    table(doc, ["Symptôme", "Cause probable", "Solution"], [
        ["« streamlit » : commande introuvable",
         "L'environnement virtuel n'est pas activé, ou PATH incomplet",
         "Activer .venv, ou lancer python -m streamlit run src/app/main.py"],
        ["ModuleNotFoundError: torch",
         "PyTorch non installé pour la plateforme",
         "pip install torch --index-url https://download.pytorch.org/whl/cpu"],
        ["ModuleNotFoundError: docx / playwright",
         "Dépendances de documentation absentes",
         "pip install -r requirements.txt (elles y figurent)"],
        ["« Aucun patient en base »",
         "La base n'a pas été initialisée",
         "Étape 5 (bouton d'accueil ou src.data.seeder)"],
        ["« Aucun modèle entraîné »",
         "Le modèle n'a jamais été entraîné sur cette source",
         "Étape 6 : page Training, onglet « Modèle final »"],
        ["participants.tsv introuvable",
         "--root ne pointe pas sur le bon dossier TDBRAIN",
         "Indiquer le dossier contenant participants.tsv (voir 8.2)"],
        ["« les données en base ne correspondent pas au modèle »",
         "Base et modèle TDBRAIN désynchronisés (contrat de features)",
         "Relancer train_tdbrain avec --seed-db pour régénérer les deux"],
        ["Le port 8501 est déjà utilisé",
         "Une autre instance tourne",
         "Ajouter --server.port 8600, ou arrêter l'instance existante"],
    ])


def _recapitulatif(doc: Document) -> None:
    heading(doc, "8. Récapitulatif des commandes")
    doc.add_paragraph("Installation minimale (cohorte simulée uniquement) :")
    code(doc, [
        "python -m venv .venv",
        ".\\.venv\\Scripts\\Activate.ps1",
        "python -m pip install --upgrade pip",
        "pip install -r requirements.txt",
        "python -m src.data.simulator",
        "python -m src.data.seeder",
        "streamlit run src/app/main.py",
        "# puis : page Training > onglet « Modèle final » > entraîner",
    ])
    doc.add_paragraph()
    doc.add_paragraph("Ajout de la cohorte réelle TDBRAIN :")
    code(doc, [
        "python -m src.models.train_tdbrain `",
        '    --root "data/tdbrain/TDBRAIN_Dataset_V3_1_Encr/TDBRAIN_Dataset_V3_1" `',
        "    --seed-db recherche_tdbrain.sqlite3",
    ])
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Une fois l'installation terminée, se reporter au « Guide d'utilisation » "
        "pour la prise en main écran par écran."
    )
    r.italic = True
    r.font.size = Pt(10)


if __name__ == "__main__":
    print(f"Wrote {build()}")
